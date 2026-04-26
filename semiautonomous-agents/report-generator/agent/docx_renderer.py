"""DOCX renderer agent.

Mirror of `renderer.py` but emits a Microsoft Word `.docx` instead of a PDF.
Runs in parallel with the PDF renderer (see `agent.py`'s ParallelAgent).

Same input contract: reads `state['report_for_render']` (a ResearchReport
JSON dict) and writes `state['docx_path']`.
"""
from __future__ import annotations

import os
from datetime import datetime, timezone
from pathlib import Path
from typing import AsyncGenerator

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from google.adk.agents import BaseAgent
from google.adk.agents.invocation_context import InvocationContext
from google.adk.events import Event
from google.genai.types import Content, Part

from .schemas import ResearchReport

_DEFAULT_OUTPUT_DIR = Path(
    os.environ.get(
        "REPORT_OUTPUT_DIR",
        Path(__file__).resolve().parents[1] / "outputs",
    )
)

_TIER_COLORS = {
    "primary": RGBColor(0x16, 0x7C, 0x3C),
    "reputable": RGBColor(0x1F, 0x4E, 0xB0),
    "secondary": RGBColor(0x88, 0x6F, 0x00),
    "unknown": RGBColor(0x66, 0x66, 0x66),
}


def _slugify(text: str) -> str:
    keep = "abcdefghijklmnopqrstuvwxyz0123456789-"
    return "".join(c if c in keep else "-" for c in text.lower()).strip("-")[:60] or "report"


def _add_horizontal_line(paragraph) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_metric(doc, metric: dict) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "EFF6FF")
    p_value = cell.paragraphs[0]
    r_value = p_value.add_run(str(metric.get("value", "")))
    r_value.font.bold = True
    r_value.font.size = Pt(20)
    r_value.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    p_label = cell.add_paragraph()
    r_label = p_label.add_run(str(metric.get("label", "")).upper())
    r_label.font.size = Pt(8)
    r_label.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    if metric.get("delta"):
        p_delta = cell.add_paragraph()
        r_delta = p_delta.add_run(str(metric["delta"]))
        r_delta.font.size = Pt(9)
        r_delta.font.bold = True
        trend = (metric.get("trend") or "flat").lower()
        r_delta.font.color.rgb = {
            "up": RGBColor(0x14, 0x53, 0x2D),
            "down": RGBColor(0x7F, 0x1D, 0x1D),
        }.get(trend, RGBColor(0x33, 0x41, 0x55))


def _add_table_block(doc, data: dict) -> None:
    headers = data.get("headers") or []
    rows = data.get("rows") or []
    if not rows:
        return
    n_cols = len(headers) if headers else max((len(r) for r in rows), default=1)
    tbl = doc.add_table(rows=(1 if headers else 0) + len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    if headers:
        hdr = tbl.rows[0].cells
        for i, h in enumerate(headers[:n_cols]):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            r = p.add_run(str(h))
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(hdr[i], "0F172A")
    start_row = 1 if headers else 0
    for ri, row in enumerate(rows):
        cells = tbl.rows[start_row + ri].cells
        for ci, cell_val in enumerate(row[:n_cols]):
            cells[ci].text = str(cell_val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if data.get("caption"):
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = cap.add_run(str(data["caption"]))
        rc.font.italic = True
        rc.font.size = Pt(9)
        rc.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def _add_comparison(doc, data: dict) -> None:
    items = data.get("items") or []
    if not items:
        return
    tbl = doc.add_table(rows=2, cols=len(items))
    tbl.autofit = True
    tbl.style = "Light Grid Accent 1"
    name_row = tbl.rows[0].cells
    body_row = tbl.rows[1].cells
    for i, item in enumerate(items):
        _shade_cell(name_row[i], "0F172A")
        p = name_row[i].paragraphs[0]
        r = p.add_run(str(item.get("name", "")))
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # body
        body_row[i].text = ""
        for h in item.get("highlights", []) or []:
            bp = body_row[i].add_paragraph(f"• {h}")
            for run in bp.runs:
                run.font.size = Pt(10)
        if item.get("verdict"):
            vp = body_row[i].add_paragraph(str(item["verdict"]))
            for run in vp.runs:
                run.font.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)


def _add_chart_placeholder(doc, data: dict) -> None:
    """DOCX doesn't get the SVG; render a labeled data table fallback."""
    title = data.get("title") or "Chart"
    p = doc.add_paragraph()
    r = p.add_run(f"Figure: {title}")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    series = data.get("series") or []
    if not series:
        return
    # collect x categories
    xs: list[str] = []
    for s in series:
        for pt in s.get("data") or []:
            x = str(pt.get("x", ""))
            if x and x not in xs:
                xs.append(x)
    if not xs:
        return
    tbl = doc.add_table(rows=1 + len(xs), cols=1 + len(series))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = data.get("x_label", "") or ""
    for i, s in enumerate(series):
        hdr[i + 1].text = s.get("name", f"Series {i + 1}")
    for ri, x in enumerate(xs):
        row = tbl.rows[1 + ri].cells
        row[0].text = x
        for si, s in enumerate(series):
            val = ""
            for pt in s.get("data") or []:
                if str(pt.get("x", "")) == x:
                    val = str(pt.get("y", ""))
                    break
            row[1 + si].text = val


def render_report_to_docx(report: ResearchReport, output_dir: Path | None = None) -> Path:
    """Public, agent-free entrypoint — useful for tests and run_local.py."""
    out_dir = Path(output_dir or _DEFAULT_OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report.topic)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x1B, 0x3D)

    if report.subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sub.add_run(report.subtitle)
        srun.font.size = Pt(14)
        srun.font.italic = True
        srun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    when = report.generated_at or datetime.now(timezone.utc).isoformat(timespec="seconds")
    try:
        when_pretty = datetime.fromisoformat(when.replace("Z", "+00:00")).strftime("%B %-d, %Y")
    except ValueError:
        when_pretty = when
    mrun = meta.add_run(f"For: {report.audience}   ·   {when_pretty}")
    mrun.font.size = Pt(10)
    mrun.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ---- Executive summary ----
    h = doc.add_heading("Executive Summary", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0F, 0x1B, 0x3D)
    doc.add_paragraph(report.executive_summary)

    if report.key_takeaways:
        kh = doc.add_heading("Key Takeaways", level=2)
        kh.runs[0].font.color.rgb = RGBColor(0x0F, 0x1B, 0x3D)
        for tk in report.key_takeaways:
            doc.add_paragraph(tk, style="List Bullet")

    doc.add_page_break()

    # ---- Sections ----
    for section in report.sections:
        sh = doc.add_heading(section.heading, level=1)
        sh.runs[0].font.color.rgb = RGBColor(0x0F, 0x1B, 0x3D)
        if section.summary:
            sp = doc.add_paragraph()
            sr = sp.add_run(section.summary)
            sr.font.italic = True
            sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for block in section.blocks:
            text = block.text or ""
            data = block.data or {}
            if block.type == "callout":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0xB0)
                _add_horizontal_line(p)
            elif block.type == "quote":
                p = doc.add_paragraph(text, style="Intense Quote")
            elif block.type == "code":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            elif block.type == "list":
                for line in text.splitlines():
                    line = line.strip().lstrip("-*•").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")
            elif block.type == "metric":
                _add_metric(doc, data)
            elif block.type == "metrics_grid":
                for m in data.get("metrics", []) or []:
                    if isinstance(m, dict):
                        _add_metric(doc, m)
            elif block.type == "table":
                _add_table_block(doc, data)
            elif block.type == "chart":
                _add_chart_placeholder(doc, data)
                if text:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = cap.add_run(text)
                    rc.font.italic = True
                    rc.font.size = Pt(9)
            elif block.type == "comparison":
                _add_comparison(doc, data)
            else:  # paragraph
                doc.add_paragraph(text)

    # ---- Open questions ----
    if report.open_questions:
        doc.add_heading("Open Questions", level=1).runs[0].font.color.rgb = RGBColor(0x0F, 0x1B, 0x3D)
        for q in report.open_questions:
            doc.add_paragraph(q, style="List Bullet")

    # ---- Sources ----
    doc.add_page_break()
    doc.add_heading("Sources", level=1).runs[0].font.color.rgb = RGBColor(0x0F, 0x1B, 0x3D)
    for idx, src in enumerate(report.sources, start=1):
        p = doc.add_paragraph()
        num = p.add_run(f"[{idx}] ")
        num.font.bold = True
        title_run = p.add_run(src.title or src.url)
        title_run.font.bold = True
        if src.domain:
            d = p.add_run(f"  ·  {src.domain}")
            d.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        tier = p.add_run(f"  [{src.credibility_tier}]")
        tier.font.color.rgb = _TIER_COLORS.get(src.credibility_tier, _TIER_COLORS["unknown"])
        tier.font.size = Pt(9)
        url_p = doc.add_paragraph()
        u = url_p.add_run(src.url)
        u.font.color.rgb = RGBColor(0x1F, 0x4E, 0xB0)
        u.font.size = Pt(9)
        if src.snippet:
            s = doc.add_paragraph()
            sr = s.add_run(src.snippet)
            sr.font.italic = True
            sr.font.size = Pt(9)
            sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    base = f"{stamp}-{_slugify(report.topic)}"
    docx_path = out_dir / f"{base}.docx"
    doc.save(str(docx_path))
    return docx_path


class DocxRendererAgent(BaseAgent):
    """Custom agent that converts state['report_for_render'] -> DOCX."""

    output_dir: Path = _DEFAULT_OUTPUT_DIR

    async def _run_async_impl(
        self, ctx: InvocationContext
    ) -> AsyncGenerator[Event, None]:
        raw = ctx.session.state.get("report_for_render") or ctx.session.state.get("report")
        if raw is None:
            yield Event(
                author=self.name,
                content=Content(parts=[Part(text="No report in state — nothing to render.")]),
            )
            return

        report = (
            raw if isinstance(raw, ResearchReport)
            else ResearchReport.model_validate(raw)
        )

        try:
            docx_path = render_report_to_docx(report, self.output_dir)
        except Exception as e:  # noqa: BLE001
            import traceback
            tb = traceback.format_exc()
            yield Event(
                author=self.name,
                content=Content(parts=[Part(text=f"DOCX rendering failed: {type(e).__name__}: {e}\n{tb}")]),
            )
            return

        ctx.session.state["docx_path"] = str(docx_path)
        yield Event(
            author=self.name,
            content=Content(parts=[Part(text=f"DOCX ready: `{docx_path}`")]),
        )
